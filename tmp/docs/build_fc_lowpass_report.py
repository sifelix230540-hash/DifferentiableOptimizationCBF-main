from __future__ import annotations

from pathlib import Path

import win32com.client as win32
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[2]
OUT_DIR = ROOT / "output" / "doc"
DOCX_PATH = OUT_DIR / "fc_layer_lowpass_report.docx"
PDF_PATH = OUT_DIR / "fc_layer_lowpass_report.pdf"


def set_east_asia_font(run, east_asia: str = "宋体") -> None:
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), east_asia)


def add_text_paragraph(
    document: Document,
    text: str,
    *,
    style: str | None = None,
    bold: bool = False,
    italic: bool = False,
    size: float = 11.5,
    align: WD_ALIGN_PARAGRAPH | None = None,
    space_after: float = 6,
) -> None:
    paragraph = document.add_paragraph(style=style)
    if align is not None:
        paragraph.alignment = align
    paragraph.paragraph_format.space_after = Pt(space_after)
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    set_east_asia_font(run)


def add_equation(document: Document, equation: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(equation)
    run.font.name = "Cambria Math"
    run.font.size = Pt(11.5)
    set_east_asia_font(run, "Cambria Math")


def add_heading(document: Document, text: str, level: int) -> None:
    heading = document.add_heading(level=level)
    heading.paragraph_format.space_after = Pt(6)
    run = heading.add_run(text)
    run.font.name = "Times New Roman"
    set_east_asia_font(run, "黑体")
    if level == 0:
        run.font.size = Pt(18)
        run.bold = True


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.start_type = WD_SECTION.NEW_PAGE

    style = document.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11.5)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def build_report() -> Document:
    doc = Document()
    configure_document(doc)

    add_heading(doc, "全连接层为何常表现为“低通滤波器”", 0)
    add_text_paragraph(
        doc,
        "副标题：从奇异值分解、图谱频率与优化动力学三条线索理解 FC 层的谱偏置",
        italic=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    add_text_paragraph(
        doc,
        "摘要：卷积层的频域解释主要来自卷积算子的拓扑约束，它在离散傅里叶基下天然近似对角化；全连接层并不存在这样的硬编码拓扑，因此它并非先天地等同于低通滤波器。更精确的说法是：在数据协方差、图拉普拉斯频率以及梯度下降/权重衰减共同作用下，训练后的 FC 权重矩阵通常会呈现低秩主导和小奇异值衰减，于是它对“全局、平滑、主方差方向”保留较多，对“局部、震荡、噪声方向”抑制更强，从而表现出数据依赖的低通性质。",
    )

    add_heading(doc, "1. 问题的严格表述", 1)
    add_text_paragraph(
        doc,
        "设单层全连接映射为 y = Wx + b，其中 x ∈ R^n，y ∈ R^m，W ∈ R^(m×n)。如果忽略偏置与非线性，FC 层就是一个一般线性算子。问题不在于它“能不能”表示低通，而在于经过训练后，它“为什么常常倾向于”保留低频/平滑模式并压制高频/噪声模式。",
    )
    add_text_paragraph(
        doc,
        "因此需要区分两个命题：",
    )
    add_text_paragraph(doc, "（1）结构命题：FC 是否因拓扑而天然等价于低通？答案是否定的。")
    add_text_paragraph(doc, "（2）训练命题：FC 在常见数据与优化条件下是否常涌现出低通行为？答案通常是肯定的。")

    add_heading(doc, "2. 代数核心：SVD 将 FC 层写成“分析-滤波-重构”", 1)
    add_text_paragraph(
        doc,
        "任意稠密矩阵 W 都可作奇异值分解：",
    )
    add_equation(doc, "W = UΣV^T = Σ(i=1 to r) σ_i u_i v_i^T,    σ_1 ≥ σ_2 ≥ ... ≥ σ_r > 0")
    add_text_paragraph(
        doc,
        "于是对输入 x 的作用可写成：",
    )
    add_equation(doc, "y = Wx = Σ(i=1 to r) σ_i u_i (v_i^T x)")
    add_text_paragraph(
        doc,
        "这个分解有一个非常清楚的滤波解释：V^T 先把输入投影到一组数据自适应基 v_i 上，Σ 再对每个分量乘以增益 σ_i，最后 U 将结果映射回输出空间。若把 v_i 看作“广义频率基”，那么奇异值 σ_i 就是对应模态的增益。",
    )
    add_text_paragraph(
        doc,
        "SVD 还给出最优低秩逼近：若只保留前 k 个奇异值，则",
    )
    add_equation(doc, "W_k = Σ(i=1 to k) σ_i u_i v_i^T")
    add_text_paragraph(
        doc,
        "并且根据 Eckart-Young 定理，有",
    )
    add_equation(doc, "||W - W_k||_F^2 = Σ(i=k+1 to r) σ_i^2")
    add_text_paragraph(
        doc,
        "这意味着：一旦训练后 W 的奇异值谱出现明显长尾，那么截去尾部小奇异值几乎不改变主要映射效果。换言之，网络主要依赖少数主导模态工作，而这些主导模态通常正对应数据中的大尺度、平滑、稳定结构。",
    )

    add_heading(doc, "3. “频率”并不一定来自傅里叶，而可以来自数据图谱", 1)
    add_text_paragraph(
        doc,
        "在卷积网络中，频率由空间平移结构决定；在全连接层中，频率更适合用“数据图”或“数据协方差”来定义。设样本图的拉普拉斯矩阵为 L = D - A，其特征分解为",
    )
    add_equation(doc, "L = ΦΛ_gΦ^T,    Λ_g = diag(λ_1^g, ..., λ_n^g),    0 = λ_1^g ≤ ... ≤ λ_n^g")
    add_text_paragraph(
        doc,
        "对图信号 x，可展开为",
    )
    add_equation(doc, "x = Σ(k=1 to n) α_k φ_k")
    add_text_paragraph(
        doc,
        "其图信号平滑度可写为",
    )
    add_equation(doc, "x^T L x = Σ(k=1 to n) λ_k^g α_k^2")
    add_text_paragraph(
        doc,
        "因此小 λ_k^g 对应平滑模态（低频），大 λ_k^g 对应剧烈变化模态（高频）。如果某个算子 T 在图谱基下满足",
    )
    add_equation(doc, "Tφ_k ≈ g(λ_k^g) φ_k,    且 |g(λ)| 随 λ 增大而减小")
    add_text_paragraph(
        doc,
        "那么 T 就是图意义下的低通滤波器。全连接层的关键区别在于：它的基不是预设的离散傅里叶基，而是通过数据与训练共同诱导出来的自适应谱基。",
    )

    add_heading(doc, "4. 为什么训练后的 FC 层会偏向低频：协方差与梯度收敛速度", 1)
    add_text_paragraph(
        doc,
        "考虑带权重衰减的线性监督学习目标：",
    )
    add_equation(doc, "L(W) = (1/2) E||Wx - y||_2^2 + (β/2)||W||_F^2")
    add_text_paragraph(
        doc,
        "设输入协方差为",
    )
    add_equation(doc, "C_x = E[xx^T] = QΛQ^T,    Λ = diag(λ_1, ..., λ_n),    λ_1 ≥ ... ≥ λ_n ≥ 0")
    add_text_paragraph(
        doc,
        "正规方程给出最优解：",
    )
    add_equation(doc, "W* = C_yx (C_x + βI)^(-1)")
    add_text_paragraph(
        doc,
        "若把问题投影到协方差特征基 q_j 上，则每个模态的解都要除以 λ_j + β。于是当 λ_j 很小（通常对应样本支持弱、方差小、噪声主导或高频细节方向）时，该方向更容易被压制。",
    )
    add_text_paragraph(
        doc,
        "更能直接看出低通性的，是自编码/去噪型目标：",
    )
    add_equation(doc, "L(W) = (1/2) E||x - Wx||_2^2 + (β/2)||W||_F^2")
    add_text_paragraph(
        doc,
        "此时闭式解为",
    )
    add_equation(doc, "W* = C_x (C_x + βI)^(-1) = Q diag( λ_j / (λ_j + β) ) Q^T")
    add_text_paragraph(
        doc,
        "这里每个模态的增益恰好是",
    )
    add_equation(doc, "g_j = λ_j / (λ_j + β)")
    add_text_paragraph(
        doc,
        "由于 g_j 随 λ_j 单调增加，大方差主模态被保留，小方差尾部模态被强烈削弱。这正是一个显式的谱域低通响应。",
    )
    add_text_paragraph(
        doc,
        "从优化动力学看，梯度下降也会优先学到大特征值方向。令单一模态上的参数为 w_j，则梯度更新近似为",
    )
    add_equation(doc, "w_j^(t+1) = (1 - η(λ_j + β)) w_j^(t) + η c_j")
    add_text_paragraph(
        doc,
        "其收敛时间尺度约为 1 / (λ_j + β)。因此 λ_j 较大的主方向更快被拟合，λ_j 较小的尾部方向不仅更慢，还更容易被正则化持续压低。这就是所谓 spectral bias 或 frequency principle 的线性化版本。",
    )

    add_heading(doc, "5. 从拓扑回看：为什么 CNN 是“硬编码低通”，FC 是“涌现低通”", 1)
    add_text_paragraph(
        doc,
        "卷积层来自局部连接与参数共享，其矩阵近似托普利茨/循环矩阵，所以在傅里叶基 F 下可近似对角化：",
    )
    add_equation(doc, "W_conv ≈ F^* diag(ĥ(ω)) F")
    add_text_paragraph(
        doc,
        "因此 CNN 的频域语言几乎是“内建”的：ĥ(ω) 就是明确的频率响应函数。",
    )
    add_text_paragraph(
        doc,
        "全连接层对应的是完全二分图 K_(n,m) 上的无约束线性映射，其矩阵既不稀疏、也不共享、也不服从固定平移对称性。它不能被先验地绑定到某个统一的傅里叶基上，只能写成",
    )
    add_equation(doc, "W_dense = UΣV^T")
    add_text_paragraph(
        doc,
        "这意味着 FC 的“频率轴”不是由空间拓扑预先规定，而是由数据分布、损失函数与训练过程共同选择出来的。CNN 的低通更像结构先验；FC 的低通更像统计规律。",
    )

    add_heading(doc, "6. 我自己的理解：FC 不是天然低通，而是天然擅长学出低秩平滑器", 1)
    add_text_paragraph(
        doc,
        "我认为把“FC 层等价于低通滤波器”直接说成绝对命题并不严谨，更严谨的表述应当是：FC 层是一个自由度极高的线性算子，它本身既可以学成低通，也可以学成高通、带通甚至近似置换；但在自然数据常见的长尾谱、有限样本、梯度下降和权重衰减共同作用下，它最容易学成的是一个低秩主导的平滑算子。",
    )
    add_text_paragraph(
        doc,
        "换句话说，FC 层的本质不是“天生低通”，而是“天生无约束”；真正把它推向低通的是统计与优化的偏置。这个结论还有两个重要推论：",
    )
    add_text_paragraph(doc, "（1）如果任务目标本身强调边缘、残差、微分或噪声放大，那么 FC 完全可能学成高频增强器。")
    add_text_paragraph(doc, "（2）如果数据协方差谱衰减很慢、没有明显主成分，或者正则化极弱，那么 FC 的低通倾向也会减弱。")
    add_text_paragraph(
        doc,
        "因此，FC 的低通性应被理解为“相对于数据流形和训练目标的低通”，而不是“相对于欧氏坐标轴的固定低通”。这也是它与卷积滤波最本质的区别。",
    )

    add_heading(doc, "7. 结论", 1)
    add_text_paragraph(
        doc,
        "可以将全文浓缩为一句话：CNN 的滤波本质来自拓扑约束导致的固定频域对角化；FC 的滤波本质来自无约束矩阵在数据谱与优化动力学下形成的奇异值塌缩和模态选择。",
    )
    add_text_paragraph(
        doc,
        "因此，若要严格描述 FC 为什么常表现为低通，最准确的三段论应当是：",
    )
    add_text_paragraph(doc, "第一，FC 层可由 SVD 写成广义谱分解，奇异值充当各模态增益。")
    add_text_paragraph(doc, "第二，数据的协方差谱或图拉普拉斯谱提供了“低频/高频”的严格定义。")
    add_text_paragraph(doc, "第三，梯度下降与权重衰减优先保留大方差、平滑、主成分方向，并抑制尾部高频方向。")
    add_text_paragraph(
        doc,
        "三者合在一起，就得到：全连接层的低通性不是由连接拓扑硬编码出来的，而是由数据谱结构与训练动力学共同涌现出来的。",
    )

    add_text_paragraph(
        doc,
        "附注：本报告为针对你的原始回答所做的“报告化”和“公式加强版”改写，因此保留了核心观点，但在“是否天然低通”这一点上增加了更严格的限定语，以避免把经验性规律误写成普适定理。",
        italic=True,
        size=10.5,
        space_after=0,
    )

    return doc


def export_pdf_via_word(docx_path: Path, pdf_path: Path) -> None:
    word = win32.DispatchEx("Word.Application")
    word.Visible = False
    document = None
    try:
        document = word.Documents.Open(str(docx_path.resolve()))
        document.ExportAsFixedFormat(
            OutputFileName=str(pdf_path.resolve()),
            ExportFormat=17,
            OpenAfterExport=False,
            OptimizeFor=0,
            CreateBookmarks=1,
        )
    finally:
        if document is not None:
            document.Close(False)
        word.Quit()


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = build_report()
    doc.save(DOCX_PATH)
    export_pdf_via_word(DOCX_PATH, PDF_PATH)
    print(f"generated: {DOCX_PATH}")
    print(f"generated: {PDF_PATH}")


if __name__ == "__main__":
    main()
